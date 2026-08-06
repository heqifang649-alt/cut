from __future__ import annotations

import argparse
import json
import math
import subprocess
from datetime import datetime, timezone
from pathlib import Path

import cv2
import numpy as np
from docx import Document


VIDEO_EXTENSIONS = {".mp4", ".mov", ".m4v", ".avi"}


def docx_text(path: Path) -> str:
    document = Document(path)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)


def audio_onsets(ffmpeg: Path, source: Path, duration: float) -> list[float]:
    seconds = min(max(duration, 1.0), 45.0)
    command = [
        str(ffmpeg), "-hide_banner", "-loglevel", "error", "-i", str(source),
        "-vn", "-t", f"{seconds:.3f}", "-ac", "1", "-ar", "1000",
        "-f", "s16le", "pipe:1",
    ]
    result = subprocess.run(command, capture_output=True, check=False)
    if result.returncode or len(result.stdout) < 2000:
        return []
    samples = np.frombuffer(result.stdout, dtype=np.int16).astype(np.float32) / 32768.0
    window = 40
    usable = len(samples) // window * window
    if usable < window * 10:
        return []
    energy = np.sqrt(np.mean(samples[:usable].reshape(-1, window) ** 2, axis=1) + 1e-9)
    rise = np.maximum(0.0, energy - np.roll(energy, 1))
    rise[0] = 0
    threshold = float(np.percentile(rise, 82))
    candidates = np.where(rise >= threshold)[0]
    onsets: list[float] = []
    for index in candidates:
        value = float(index * window / 1000.0)
        if not onsets or value - onsets[-1] >= 0.18:
            onsets.append(value)
        elif rise[index] > rise[int(round(onsets[-1] * 1000 / window))]:
            onsets[-1] = value
    return onsets


def frame_lab(frame: np.ndarray) -> np.ndarray:
    small = cv2.resize(frame, (96, 170), interpolation=cv2.INTER_AREA)
    return cv2.cvtColor(small, cv2.COLOR_BGR2LAB).reshape(-1, 3).mean(axis=0)


def analyze_video(source: Path, ffmpeg: Path, contact_sheet: Path) -> dict:
    capture = cv2.VideoCapture(str(source))
    fps = float(capture.get(cv2.CAP_PROP_FPS) or 30.0)
    frame_count = int(capture.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    width = int(capture.get(cv2.CAP_PROP_FRAME_WIDTH) or 0)
    height = int(capture.get(cv2.CAP_PROP_FRAME_HEIGHT) or 0)
    duration = frame_count / fps if fps > 0 else 0.0
    sample_fps = min(8.0, fps)
    step = max(1, int(round(fps / sample_fps)))
    max_frames = int(min(frame_count, max(1, math.ceil(min(duration, 45.0) * fps))))

    labs: list[np.ndarray] = []
    histograms: list[np.ndarray] = []
    timestamps: list[float] = []
    cut_candidates: list[tuple[float, float]] = []
    motion_jerks: list[float] = []
    sharpness: list[float] = []
    previous_gray = None
    previous_translation = None
    previous_hist = None
    previous_lab = None
    index = 0
    while index < max_frames:
        capture.set(cv2.CAP_PROP_POS_FRAMES, index)
        ok, frame = capture.read()
        if not ok:
            break
        timestamp = index / fps
        gray = cv2.cvtColor(cv2.resize(frame, (192, 340), interpolation=cv2.INTER_AREA), cv2.COLOR_BGR2GRAY)
        lab = frame_lab(frame)
        hsv = cv2.cvtColor(cv2.resize(frame, (96, 170), interpolation=cv2.INTER_AREA), cv2.COLOR_BGR2HSV)
        histogram = cv2.calcHist([hsv], [0, 1], None, [24, 24], [0, 180, 0, 256])
        cv2.normalize(histogram, histogram)
        sharpness.append(float(cv2.Laplacian(gray, cv2.CV_64F).var()))

        if previous_hist is not None and previous_lab is not None:
            correlation = float(cv2.compareHist(previous_hist, histogram, cv2.HISTCMP_CORREL))
            lab_delta = float(np.linalg.norm(lab - previous_lab))
            score = max(0.0, 1.0 - correlation) * 0.7 + min(lab_delta / 60.0, 1.0) * 0.3
            if score > 0.48:
                cut_candidates.append((timestamp, score))

        if previous_gray is not None:
            points = cv2.goodFeaturesToTrack(previous_gray, maxCorners=120, qualityLevel=0.02, minDistance=7)
            if points is not None and len(points) >= 10:
                moved, status, _ = cv2.calcOpticalFlowPyrLK(previous_gray, gray, points, None)
                valid = status.reshape(-1) == 1
                if valid.sum() >= 8:
                    delta = moved[valid] - points[valid]
                    translation = np.median(delta.reshape(-1, 2), axis=0)
                    normalized = translation / np.array([gray.shape[1], gray.shape[0]], dtype=np.float32)
                    if previous_translation is not None:
                        motion_jerks.append(float(np.linalg.norm(normalized - previous_translation)))
                    previous_translation = normalized

        labs.append(lab)
        histograms.append(histogram)
        timestamps.append(timestamp)
        previous_gray = gray
        previous_hist = histogram
        previous_lab = lab
        index += step
    capture.release()

    cuts: list[float] = []
    for timestamp, score in sorted(cut_candidates, key=lambda item: item[1], reverse=True):
        if all(abs(timestamp - existing) >= 0.45 for existing in cuts):
            cuts.append(timestamp)
    cuts.sort()
    cuts = [value for value in cuts if 0.25 < value < max(0.25, duration - 0.25)]

    onsets = audio_onsets(ffmpeg, source, duration)
    errors = [min((abs(cut - onset) for onset in onsets), default=1.0) for cut in cuts]
    lab_array = np.asarray(labs) if labs else np.zeros((1, 3))
    color_deltas = np.linalg.norm(np.diff(lab_array, axis=0), axis=1) if len(lab_array) > 1 else np.array([0.0])
    jerk_array = np.asarray(motion_jerks) if motion_jerks else np.array([0.0])

    sheet_frames = []
    capture = cv2.VideoCapture(str(source))
    for ratio in np.linspace(0.04, 0.96, 9):
        capture.set(cv2.CAP_PROP_POS_MSEC, duration * ratio * 1000)
        ok, frame = capture.read()
        if not ok:
            continue
        thumb = cv2.resize(frame, (180, 320), interpolation=cv2.INTER_AREA)
        cv2.putText(thumb, f"{duration * ratio:.1f}s", (8, 24), cv2.FONT_HERSHEY_SIMPLEX, 0.55, (255, 255, 255), 2, cv2.LINE_AA)
        sheet_frames.append(thumb)
    capture.release()
    if sheet_frames:
        rows = []
        for start in range(0, len(sheet_frames), 3):
            row = sheet_frames[start:start + 3]
            while len(row) < 3:
                row.append(np.zeros_like(sheet_frames[0]))
            rows.append(np.hstack(row))
        encoded, buffer = cv2.imencode(".jpg", np.vstack(rows))
        if encoded:
            buffer.tofile(str(contact_sheet))

    shot_count = len(cuts) + 1
    return {
        "file": source.name,
        "duration_seconds": round(duration, 3),
        "width": width,
        "height": height,
        "fps": round(fps, 3),
        "aspect_ratio": round(width / height, 4) if height else None,
        "detected_cut_times": [round(value, 3) for value in cuts],
        "detected_shot_count": shot_count,
        "average_shot_length_seconds": round(duration / shot_count, 3) if shot_count else None,
        "first_cut_seconds": round(cuts[0], 3) if cuts else None,
        "color": {
            "mean_lab": [round(float(value), 3) for value in lab_array.mean(axis=0)],
            "sample_delta_median": round(float(np.median(color_deltas)), 3),
            "sample_delta_p90": round(float(np.percentile(color_deltas, 90)), 3),
        },
        "stability": {
            "camera_jerk_median": round(float(np.median(jerk_array)), 6),
            "camera_jerk_p90": round(float(np.percentile(jerk_array, 90)), 6),
        },
        "sharpness_median": round(float(np.median(sharpness)), 3) if sharpness else None,
        "audio": {
            "detected_onsets": [round(value, 3) for value in onsets[:80]],
            "cut_to_beat_error_median_seconds": round(float(np.median(errors)), 3) if errors else None,
            "cut_to_beat_error_p90_seconds": round(float(np.percentile(errors, 90)), 3) if errors else None,
        },
        "contact_sheet": contact_sheet.name,
    }


def percentile(values: list[float], amount: float, fallback: float) -> float:
    return round(float(np.percentile(values, amount)), 4) if values else fallback


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--videos", required=True, type=Path)
    parser.add_argument("--document", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--ffmpeg", required=True, type=Path)
    args = parser.parse_args()

    args.output.mkdir(parents=True, exist_ok=True)
    profiles_dir = args.output / "profiles"
    sheets_dir = args.output / "contact-sheets"
    profiles_dir.mkdir(exist_ok=True)
    sheets_dir.mkdir(exist_ok=True)
    videos = sorted(path for path in args.videos.rglob("*") if path.is_file() and path.suffix.lower() in VIDEO_EXTENSIONS)
    profiles = []
    for video in videos:
        safe_name = f"{len(profiles) + 1:02d}-{video.stem}"
        profile = analyze_video(video, args.ffmpeg, sheets_dir / f"{safe_name}.jpg")
        profiles.append(profile)
        (profiles_dir / f"{safe_name}.json").write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")

    shot_lengths = [item["average_shot_length_seconds"] for item in profiles if item["average_shot_length_seconds"]]
    color_p90 = [item["color"]["sample_delta_p90"] for item in profiles]
    jitter_p90 = [item["stability"]["camera_jerk_p90"] for item in profiles]
    beat_errors = [item["audio"]["cut_to_beat_error_median_seconds"] for item in profiles if item["audio"]["cut_to_beat_error_median_seconds"] is not None]
    clusters = {
        "under_13_seconds": [item["file"] for item in profiles if item["duration_seconds"] <= 13.0],
        "13_to_20_seconds": [item["file"] for item in profiles if 13.0 < item["duration_seconds"] <= 20.0],
        "over_20_seconds": [item["file"] for item in profiles if item["duration_seconds"] > 20.0],
    }
    standard = {
        "schema_version": "gc-fashion-gold-standard/1.0",
        "dataset_id": "gc-good-video-set-20260805",
        "version": 1,
        "status": "approved",
        "approved_at": datetime.now(timezone.utc).isoformat(),
        "scope": "海外信息流服装广告，Cutflow 13秒竖屏批量初剪",
        "sources": {
            "reference_directory": str(args.videos),
            "reference_document": str(args.document),
            "video_count": len(profiles),
            "performance_data_available": False,
            "claim_limit": "仅定义创意与制作质量，不宣称已证明转化效果。",
        },
        "watch_logic": ["停留", "兴趣", "价值", "转化"],
        "hard_gates": [
            {"id": "same_product", "rule": "全片只出现同一款服装；不确定则排除。"},
            {"id": "original_speed", "rule": "模特动作保持1.00倍，不通过加速制造节奏。"},
            {"id": "product_by_one_second", "rule": "1秒内清楚看到服装或核心产品信息。"},
            {"id": "stable_motion", "rule": "画面运镜稳定，无高频抖动、对焦漂移和无目的晃动。"},
            {"id": "color_continuity", "rule": "前后镜头白平衡、黑位、饱和度和氛围统一，无明显色调跳变。"},
            {"id": "shot_coverage", "rule": "同时包含整体上身展示和至少一个衣服细节镜头，聚焦衣服本身。"},
            {"id": "beat_sync", "rule": "主要切点贴合BGM强拍或能量上升点，不牺牲动作完整性。"},
            {"id": "short_video_logic", "rule": "结构遵循停留、兴趣、价值、转化，禁止长铺垫和低信息密度。"},
            {"id": "technical", "rule": "不超过13秒、可完整解码、字幕在安全区、核心图案无遮挡。"},
        ],
        "cutflow_structure": [
            {"time": "0-3s", "role": "停留", "requirement": "最强服装识别或动作Hook，产品最好在1秒内出现。"},
            {"time": "3-6s", "role": "兴趣", "requirement": "完整上身、版型和穿搭关系。"},
            {"time": "6-8.2s", "role": "价值1", "requirement": "正面图案、领口或版型。"},
            {"time": "8.2-10.2s", "role": "价值2", "requirement": "袖口、面料、刺绣或工艺。"},
            {"time": "10.2-12.7s", "role": "转化", "requirement": "背面或最佳补充镜头，稳定收口并留CVR阅读时间。"},
        ],
        "weights": {
            "product_focus_and_coverage": 25,
            "rhythm_and_short_video_logic": 22,
            "color_atmosphere_and_continuity": 18,
            "camera_stability": 15,
            "music_and_beat_sync": 12,
            "technical_and_safe_zone": 8,
        },
        "reference_statistics": {
            "average_shot_length_median_seconds": percentile(shot_lengths, 50, 1.5),
            "average_shot_length_p75_seconds": percentile(shot_lengths, 75, 2.2),
            "color_sample_delta_p90_median": percentile(color_p90, 50, 12.0),
            "camera_jerk_p90_p75": percentile(jitter_p90, 75, 0.018),
            "cut_to_beat_error_median_p75_seconds": percentile(beat_errors, 75, 0.18),
        },
        "automatic_review": {
            "minimum_total_score": 95,
            "hard_gate_failure_blocks_review": True,
            "only_failed_product_is_revised": True,
        },
        "explicit_rejections": [
            "开头人物走路、看镜头、摆姿势等铺垫超过1秒但产品不清楚",
            "单镜头长时间没有新信息",
            "同一成片混入多件衣服",
            "镜头抖动、动作跳断、色温或饱和度明显跳变",
            "切点脱离BGM且通过改变动作速度强行卡点",
            "只有整体没有细节，或只有细节无法确认上身效果",
            "字幕、手势或CVR遮挡衣服核心图案",
        ],
        "clusters": clusters,
        "profiles": [item["file"] for item in profiles],
        "document_excerpt": docx_text(args.document),
    }
    (args.output / "clusters.json").write_text(json.dumps(clusters, ensure_ascii=False, indent=2), encoding="utf-8")
    (args.output / "gold-standard-v1.json").write_text(json.dumps(standard, ensure_ascii=False, indent=2), encoding="utf-8")
    (args.output / "review.md").write_text(
        "# GC优秀视频标准 v1\n\n"
        f"- 状态：已确认\n- 样片数：{len(profiles)}\n- 适用范围：海外信息流服装广告，13秒竖屏批量初剪\n\n"
        "## 核心共识\n\n"
        "色调前后一致且有氛围；运镜稳定；切点与BGM配合；整体和细节均有覆盖；画面持续聚焦服装；结构遵循停留、兴趣、价值、转化。\n",
        encoding="utf-8",
    )
    print(json.dumps({"videos": len(profiles), "output": str(args.output), "standard": str(args.output / 'gold-standard-v1.json')}, ensure_ascii=False))


if __name__ == "__main__":
    main()
